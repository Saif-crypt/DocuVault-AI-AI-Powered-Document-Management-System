"""
ocr_engine.py
─────────────────────────────────────────────────────────────────────────────
AI-Powered Document Management System — OCR Module
─────────────────────────────────────────────────────────────────────────────
Purpose   : Extracts text from scanned images and scanned PDFs using Tesseract.
Usage     : python ocr_engine.py <file_path>
Output    : Extracted text printed to stdout, wrapped in START_TEXT / END_TEXT markers.
            Nothing else is printed to stdout. Errors go to stderr.
─────────────────────────────────────────────────────────────────────────────
Design Notes (for future integration):
  • No NLP, ML, database, web-framework, or cloud dependencies.
  • Every processing stage is isolated in its own function → easy to extend.
  • PDF→image conversion, preprocessing, and OCR are each separate concerns.
─────────────────────────────────────────────────────────────────────────────
"""

import sys
import os
import io
import cv2
import numpy as np
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from docx import Document  # Add this line to import python-docx
# Add this line to specify Tesseract location
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


# Force UTF-8 encoding
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except AttributeError:
    pass  # Python < 3.7

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────

# File extensions the module understands
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"}
SUPPORTED_PDF_EXTENSION   = ".pdf"
SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx"}  # Add this new line

# Markers that wrap the final output (required by ASP.NET integration spec)
OUTPUT_START_MARKER = "-----START_TEXT-----"
OUTPUT_END_MARKER   = "-----END_TEXT-----"


# ─────────────────────────────────────────────────────────────────────────────
# 1. INPUT VALIDATION
# ─────────────────────────────────────────────────────────────────────────────

def validate_input(file_path: str) -> str:
    """
    Validates the given file path.
    Returns the resolved absolute path on success.
    Raises FileNotFoundError or ValueError on failure.
    """
    file_path = os.path.abspath(file_path)

    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = os.path.splitext(file_path)[1].lower()
    all_supported = SUPPORTED_IMAGE_EXTENSIONS | SUPPORTED_DOCUMENT_EXTENSIONS

    if extension not in all_supported:
        raise ValueError(
            f"Unsupported file format: '{extension}'. "
            f"Supported formats: {', '.join(sorted(all_supported))}"
        )

    return file_path


# ─────────────────────────────────────────────────────────────────────────────
# 2. FILE-TYPE DETECTION
# ─────────────────────────────────────────────────────────────────────────────

def detect_file_type(file_path: str) -> str:
    """
    Returns 'image', 'pdf', or 'docx' based on the file extension.
    Assumes validate_input() was already called.
    """
    extension = os.path.splitext(file_path)[1].lower()
    if extension == SUPPORTED_PDF_EXTENSION:
        return "pdf"
    elif extension == ".docx":
        return "docx"
    return "image"


# ─────────────────────────────────────────────────────────────────────────────
# 3. PDF → IMAGE CONVERSION
# ─────────────────────────────────────────────────────────────────────────────

def pdf_to_images(pdf_path: str) -> list:
    """
    Converts every page of a PDF into a PIL Image object.
    Uses pdf2image (which wraps Poppler internally) at 300 DPI.

    Returns a list of PIL.Image objects, one per page.
    Raises RuntimeError if conversion fails.
    """
    try:
        from pdf2image import convert_from_path

        images = convert_from_path(
            pdf_path,
            dpi=300
        )
        return images

    except ImportError:
        raise RuntimeError(
            "pdf2image is not installed. Run:  pip install pdf2image"
        )
    except Exception as e:
        raise RuntimeError(
            f"PDF conversion failed: {e}\n"
            f"Make sure Poppler is installed and added to your PATH."
        )


# ─────────────────────────────────────────────────────────────────────────────
# 4. IMAGE PREPROCESSING
# ─────────────────────────────────────────────────────────────────────────────

def preprocess_image(pil_image: Image.Image) -> Image.Image:
    """
    Applies a preprocessing pipeline that typically improves OCR accuracy
    on scanned documents:

        1. Convert to grayscale
        2. Resize (upscale small images to >= 300 DPI equivalent)
        3. Adaptive thresholding → clean black-and-white text

    Accepts a PIL Image, returns a new PIL Image (grayscale, thresholded).
    The original image is NOT modified.
    """
    # --- Step 1: Convert PIL → OpenCV (numpy array) ---
    img_array = np.array(pil_image)

    # --- Step 2: Grayscale conversion ---
    if len(img_array.shape) == 3:
        gray = cv2.cvtColor(img_array, cv2.COLOR_BGR2GRAY)
    else:
        gray = img_array

    # --- Step 3: Upscale if image is too small ---
    height, width = gray.shape[:2]
    if height < 1500 or width < 1500:
        gray = cv2.resize(gray, None, fx=2, fy=2, interpolation=cv2.INTER_LANCZOS4)

    # --- Step 4: Adaptive thresholding ---
    # IMPORTANT: OpenCV requires ALL arguments to be POSITIONAL here.
    #            Using keyword arguments causes a crash on Windows builds.
    #
    # Argument order:
    #   1. src             → grayscale input image
    #   2. maxValue        → 255 (white)
    #   3. adaptiveMethod  → GAUSSIAN_C (handles uneven lighting well)
    #   4. thresholdType   → THRESH_BINARY (output is black & white)
    #   5. blockSize       → 15 (neighbourhood size, must be odd)
    #   6. C               → 10 (constant subtracted from mean)
    binary = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, 10)

    # --- Step 5: Convert back to PIL Image for Tesseract ---
    return Image.fromarray(binary)

def extract_text_from_docx(docx_path):
    """
    Extract text from a DOCX file.
    
    Args:
        docx_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from the document
    """
    try:
        doc = Document(docx_path)
        
        # Extract text from paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")

# ─────────────────────────────────────────────────────────────────────────────
# 5. OCR EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_image(pil_image: Image.Image) -> str:
    """
    Runs Tesseract OCR on a single preprocessed PIL Image.
    Returns the raw extracted text string.

    Tesseract config flags used:
        --oem 3  : Use the best available OCR engine (LSTM when available)
        --psm 3  : Fully automatic page segmentation (default, but explicit)
    """
    preprocessed = preprocess_image(pil_image)

    text = pytesseract.image_to_string(
        preprocessed,
        lang="eng",
        config="--oem 3 --psm 3"
    )
    return text


# ─────────────────────────────────────────────────────────────────────────────
# 6. MAIN PIPELINE
# ─────────────────────────────────────────────────────────────────────────────

def run_ocr(file_path: str) -> str:
    """
    Full pipeline:
        validate → detect type → (convert PDF if needed) → preprocess → OCR
    Returns the combined extracted text from all pages / images.
    For DOCX files, extracts text directly without OCR.
    """
    file_path = validate_input(file_path)
    file_type = detect_file_type(file_path)

    # Handle DOCX files - direct text extraction (no OCR needed)
    if file_type == "docx":
        return extract_text_from_docx(file_path)
    
    # Handle PDF and image files with OCR
    if file_type == "pdf":
        images = pdf_to_images(file_path)
    else:
        images = [Image.open(file_path)]

    all_text_parts = []
    for image in images:
        page_text = extract_text_from_image(image)
        all_text_parts.append(page_text)
        image.close()

    combined_text = "\n\n".join(all_text_parts)
    return combined_text


# ─────────────────────────────────────────────────────────────────────────────
# 7. ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def main():
    """
    CLI entry-point.
    Reads file path from sys.argv[1], runs OCR, and prints the result
    wrapped in the mandatory START_TEXT / END_TEXT markers.

    ALL error messages go to stderr so stdout stays clean for the ASP.NET parser.
    """
    if len(sys.argv) < 2:
        print(
            "Usage: python ocr_engine.py <file_path>\n"
            "Example: python ocr_engine.py test_input/sample_scan.png",
            file=sys.stderr
        )
        sys.exit(1)

    input_path = sys.argv[1]

    try:
        extracted_text = run_ocr(input_path)

        # Print ONLY the wrapped text to stdout — nothing else
        print(OUTPUT_START_MARKER)
        print(extracted_text)
        print(OUTPUT_END_MARKER)

    except FileNotFoundError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)

    except ValueError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)

    except RuntimeError as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)

    except Exception as e:
        print(f"[ERROR] Unexpected failure: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
